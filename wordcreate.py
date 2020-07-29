from docx import Document

d = Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Hello this is another paragraph.')
p=d.paragraphs[0]
p.add_run('This is a new run..!')
p.runs[1].bold=True

d.save('F:\\Python\\XL-Word-PDF\\demo4.docx')
